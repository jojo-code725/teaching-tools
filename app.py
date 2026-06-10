"""
教培工具网页版 v2 — 学生档案生成 + 课后反馈生成
支持在线填写 & Excel上传两种模式
运行: python3 -m streamlit run app.py
"""

import streamlit as st
import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from collections import defaultdict
import io, json, os, tempfile, zipfile, copy, re, ssl
import urllib.request, urllib.error
ssl._create_default_https_context = ssl._create_unverified_context

st.set_page_config(page_title="教培效率工具", page_icon="📚", layout="wide")
st.title("📚 教培效率工具 v2")

# ============================================================
# 共享工具函数（文档生成）
# ============================================================
PAGE_W = Cm(29.7); PAGE_H = Cm(21.0)
MARGIN_LR = Cm(2.54); MARGIN_TB = Cm(3.05)
TABLE_HEADERS = ["序号","科目","托管老师","月考成绩","班级排名","总分","学校排名",
                 "月度学习表现","本月考情分析","下一阶段学科建议","备注"]

def set_page(doc):
    sec = doc.sections[0]
    sec.page_width = PAGE_W; sec.page_height = PAGE_H
    sec.left_margin = MARGIN_LR; sec.right_margin = MARGIN_LR
    sec.top_margin = MARGIN_TB; sec.bottom_margin = MARGIN_TB

def add_title(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("一鸣学员档案"); run.font.size = Pt(14); run.bold = True

def add_info_line(doc, school, grade, name, enroll):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    text = f"学校：{school}    年级班级：{grade}         姓名：{name}       入学时间：{enroll}"
    run = p.add_run(text); run.font.size = Pt(14); run.bold = True

def add_subject_header(doc, name):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(name); run.font.size = Pt(12); run.bold = True

def add_label(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); run.font.size = Pt(12); run.bold = True

def add_dual(doc, label, text):
    p = doc.add_paragraph()
    r1 = p.add_run(label); r1.font.size = Pt(12); r1.bold = True
    if text: p.add_run(text)

def add_body(doc, text):
    if not text: return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.add_run(text)

def add_body_lines(doc, text):
    if not text: return
    for line in text.replace("\r","").split("\n"):
        line = line.strip()
        if not line: continue
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; p.add_run(line)

def add_summary_table(doc, subjects):
    n = len(subjects)
    table = doc.add_table(rows=n+1, cols=11)
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(TABLE_HEADERS):
        cell = table.rows[0].cells[j]; cell.text = ""
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h); run.font.size = Pt(9); run.bold = True
    for i, subj in enumerate(subjects):
        row = table.rows[i+1]
        vals = [str(i+1), subj.get("科目",""), subj.get("托管老师",""),
                subj.get("月考成绩",""), subj.get("班级排名",""),
                subj.get("总分",""), subj.get("学校排名",""), "","","",""]
        for j, v in enumerate(vals):
            cell = row.cells[j]; cell.text = ""
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run(v)

def build_archive_doc(name, info, subjects):
    """生成单个学生档案 docx"""
    doc = Document(); set_page(doc)
    add_title(doc)
    add_info_line(doc, info.get("学校",""), info.get("年级班级",""), name, info.get("入学时间",""))
    doc.add_paragraph()
    for subj in subjects:
        add_subject_header(doc, subj["科目"])
        add_label(doc, "月度学习表现：")
        add_dual(doc, "优点：", subj.get("优点",""))
        add_dual(doc, "不足：", subj.get("不足",""))
        add_label(doc, "考情分析：")
        add_body(doc, subj.get("考情分析",""))
        add_label(doc, "下一阶段建议：")
        add_body_lines(doc, subj.get("下一阶段建议",""))
        doc.add_paragraph()
    add_summary_table(doc, subjects)
    return doc

# ============================================================
# AI 文件总结
# ============================================================
def extract_text(file_bytes, filename):
    """从上传文件中提取文本"""
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.txt':
        return file_bytes.decode('utf-8', errors='ignore')
    elif ext == '.docx':
        from docx import Document as DocRead
        from io import BytesIO
        doc = DocRead(BytesIO(file_bytes))
        return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
    elif ext == '.pdf':
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_bytes))
            return '\n'.join([page.extract_text() or '' for page in reader.pages])
        except ImportError:
            return "[PDF读取需要安装PyPDF2，请在requirements.txt中添加]"
    else:
        return f"[不支持的文件格式: {ext}]"

def ai_summarize(text, api_key, max_chars=300):
    """调用DeepSeek API总结文本"""
    if not api_key:
        return None, "请先输入API Key"
    if len(text) < 50:
        return None, "文本太短，无需总结"

    # 截取前8000字发送给AI
    prompt_text = text[:8000]

    payload = json.dumps({
        "model": "deepseek-chat",
        "messages": [
            {"role": "system", "content": "你是一个专业的教育内容总结助手。请将以下教学内容总结为约300字的上课内容概述，用中文输出，分为2-3段，突出知识点、教学重点和课堂练习内容。直接输出总结，不要加前缀说明。"},
            {"role": "user", "content": f"请总结以下教学内容：\n\n{prompt_text}"}
        ],
        "max_tokens": 600,
        "temperature": 0.5
    }).encode('utf-8')

    req = urllib.request.Request(
        "https://api.deepseek.com/v1/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {api_key}"
        }
    )
    try:
        with urllib.request.urlopen(req, timeout=60) as resp:
            result = json.loads(resp.read())
            return result["choices"][0]["message"]["content"].strip(), None
    except urllib.error.HTTPError as e:
        err = json.loads(e.read())
        return None, f"API错误: {err.get('message', str(e))}"
    except Exception as e:
        return None, f"请求失败: {str(e)}"

# ============================================================
# 初始化 session_state
# ============================================================
if "archive_students" not in st.session_state:
    st.session_state.archive_students = {}  # {name: {info: {}, subjects: []}}
if "fb_students" not in st.session_state:
    st.session_state.fb_students = []  # [{name, notes}]

# ============================================================
# Tab 1: 学生档案生成
# ============================================================
tab1, tab2 = st.tabs(["📋 学生档案生成", "✏️ 课后反馈生成"])

with tab1:
    st.header("学生档案生成")

    mode = st.radio("选择输入方式", ["✏️ 在线填写", "📊 上传Excel"], horizontal=True, key="archive_mode")

    if mode == "📊 上传Excel":
        col1, col2 = st.columns(2)
        with col1:
            excel_file = st.file_uploader("学生更新数据 Excel", type=["xlsx"], key="excel_up")
        with col2:
            json_file = st.file_uploader("学生基础信息库 JSON", type=["json"], key="json_up")

        if excel_file and json_file:
            try:
                base_info = json.loads(json_file.read())
                wb = openpyxl.load_workbook(excel_file)
                ws = wb["学生更新数据"]
                students = defaultdict(lambda: {"info": {}, "subjects": []})
                for row in ws.iter_rows(min_row=2, values_only=True):
                    if not row or not row[0]: continue
                    vals = [str(c).strip() if c else "" for c in row]
                    name = vals[0]
                    if not name: continue
                    base = base_info.get(name, {})
                    teacher = ""
                    for bs in base.get("科目",[]):
                        if bs["科目"] == vals[1]: teacher = bs["托管老师"]; break
                    if not students[name]["info"]:
                        students[name]["info"] = {"学校": base.get("学校",""),
                            "年级班级": base.get("年级班级",""), "入学时间": base.get("入学时间","")}
                    students[name]["subjects"].append({
                        "科目": vals[1], "托管老师": teacher,
                        "月考成绩": vals[2], "班级排名": vals[3],
                        "总分": vals[4], "学校排名": vals[5],
                        "优点": vals[6], "不足": vals[7],
                        "考情分析": vals[8], "下一阶段建议": vals[9]})

                if st.button("🚀 批量生成档案", type="primary", key="gen_xlsx"):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        zip_path = os.path.join(tmpdir, "档案汇总.zip")
                        with zipfile.ZipFile(zip_path, 'w') as zf:
                            for name, data in students.items():
                                doc = build_archive_doc(name, data["info"], data["subjects"])
                                safe = name.replace("/","_")
                                dp = os.path.join(tmpdir, f"{safe}.docx")
                                doc.save(dp); zf.write(dp, f"{safe}.docx")
                        with open(zip_path, "rb") as f:
                            st.download_button("📥 下载全部档案（ZIP）", f.read(),
                                               "学生档案汇总.zip", "application/zip")
                        st.success(f"✅ 已生成 {len(students)} 份档案")
            except Exception as e:
                st.error(f"处理失败：{e}")

    else:
        # ===== 在线填写模式 =====
        # 加载基础信息库
        base_json = st.file_uploader("📂 加载基础信息库（可选，用于自动补全）", type=["json"], key="base_json_online")
        base_info = {}
        if base_json:
            base_info = json.loads(base_json.read())
            st.success(f"已加载 {len(base_info)} 名学生的基础信息")

        st.divider()

        # --- 学生基本信息 ---
        st.subheader("👤 学生信息")

        # 恢复草稿
        draft_file = st.file_uploader("📂 恢复草稿（上传之前保存的 .json 文件）", type=["json"], key="draft_load")
        restored_name = ""
        if draft_file:
            draft = json.loads(draft_file.read())
            st.session_state._restored_draft = draft
        if "_restored_draft" in st.session_state:
            draft = st.session_state._restored_draft
            restored_name = draft.get("_student_name", "")
            if f"subj_count_{restored_name}" not in st.session_state:
                st.session_state[f"subj_count_{restored_name}"] = draft.get("subj_count", 4)

        c1, c2 = st.columns([1, 3])
        with c1:
            if base_info:
                name_sel = st.selectbox("快速选择", ["--新学生--"] + list(base_info.keys()), key="name_sel")
            else:
                name_sel = "--新学生--"
                st.caption("加载基础库后可快速选择")
        with c2:
            manual_name = st.text_input("姓名", value=restored_name if restored_name else ("" if name_sel == "--新学生--" else name_sel), key="manual_name")

        st_name = manual_name or (name_sel if name_sel != "--新学生--" else "")

        # 自动补全
        auto_info = base_info.get(st_name, {}) if st_name else {}
        # 草稿数据优先
        draft_info = {}
        if "_restored_draft" in st.session_state:
            d = st.session_state._restored_draft
            draft_info = {"学校": d.get("school",""), "年级班级": d.get("grade",""), "入学时间": d.get("enroll","")}
        c3, c4, c5 = st.columns(3)
        with c3:
            school = st.text_input("学校", value=draft_info.get("学校") or auto_info.get("学校",""), key="school")
        with c4:
            grade = st.text_input("年级班级", value=draft_info.get("年级班级") or auto_info.get("年级班级",""), key="grade")
        with c5:
            enroll = st.text_input("入学时间", value=draft_info.get("入学时间") or auto_info.get("入学时间",""), key="enroll")

        default_subs = auto_info.get("科目", [{"科目":"英语","托管老师":""},{"科目":"数学","托管老师":""},{"科目":"物理","托管老师":""},{"科目":"语文","托管老师":""}])

        # --- 成绩汇总表 ---
        st.subheader("📊 成绩汇总")
        if f"subj_count_{st_name}" not in st.session_state:
            st.session_state[f"subj_count_{st_name}"] = max(len(default_subs), 1)

        subj_count = st.session_state[f"subj_count_{st_name}"]
        subj_data = []

        # Table header
        h_cols = st.columns([2, 1.5, 1, 1, 1, 1])
        headers = ["科目", "老师", "成绩", "班排", "校排", "总分"]
        for hc, h in zip(h_cols, headers):
            with hc: st.caption(f"**{h}**")

        # Score rows - compact
        draft_subs = st.session_state.get("_restored_draft", {}).get("subj_data", [])
        for i in range(subj_count):
            ds = default_subs[i] if i < len(default_subs) else {"科目":"", "托管老师":""}
            dd = draft_subs[i] if i < len(draft_subs) else {}
            c1, c2, c3, c4, c5, c6 = st.columns([2, 1.5, 1, 1, 1, 1])
            with c1:
                subj_name = st.text_input("科目", value=dd.get("科目") or ds.get("科目",""), key=f"sn_{st_name}_{i}", label_visibility="collapsed", placeholder="科目")
            with c2:
                teacher = st.text_input("老师", value=dd.get("托管老师") or ds.get("托管老师",""), key=f"st_{st_name}_{i}", label_visibility="collapsed", placeholder="老师")
            with c3:
                score = st.text_input("成绩", value=dd.get("月考成绩",""), key=f"ss_{st_name}_{i}", label_visibility="collapsed", placeholder="0")
            with c4:
                class_rank = st.text_input("班排", value=dd.get("班级排名",""), key=f"scr_{st_name}_{i}", label_visibility="collapsed", placeholder="0")
            with c5:
                school_rank = st.text_input("校排", value=dd.get("学校排名",""), key=f"ssr_{st_name}_{i}", label_visibility="collapsed", placeholder="0")
            with c6:
                total_score = st.text_input("总分", value=dd.get("总分",""), key=f"sts_{st_name}_{i}", label_visibility="collapsed", placeholder="0")

            subj_data.append({
                "科目": subj_name, "托管老师": teacher,
                "月考成绩": score, "班级排名": class_rank,
                "总分": total_score, "学校排名": school_rank,
                "优点": "", "不足": "", "考情分析": "", "下一阶段建议": ""
            })

        col_add, col_del, _ = st.columns([1, 1, 4])
        with col_add:
            if st.button("➕ 加一科", key="add_subj"):
                st.session_state[f"subj_count_{st_name}"] += 1; st.rerun()
        with col_del:
            if st.button("➖ 减一科", key="del_subj") and st.session_state[f"subj_count_{st_name}"] > 1:
                st.session_state[f"subj_count_{st_name}"] -= 1; st.rerun()

        # --- 科目评语 ---
        st.divider()
        st.subheader("📝 科目评语")

        for i in range(subj_count):
            sn = subj_data[i]["科目"] if i < len(subj_data) else f"科目{i+1}"
            dd_comment = draft_subs[i] if i < len(draft_subs) else {}
            with st.expander(f"{sn or '新科目'}", expanded=(i==0)):
                adv = st.text_area("优点", value=dd_comment.get("优点",""), key=f"sa_{st_name}_{i}", height=68)
                weak = st.text_area("不足", value=dd_comment.get("不足",""), key=f"sw_{st_name}_{i}", height=68)
                exam = st.text_area("考情分析", value=dd_comment.get("考情分析",""), key=f"se_{st_name}_{i}", height=100)
                suggest = st.text_area("下一阶段建议", value=dd_comment.get("下一阶段建议",""), key=f"ssg_{st_name}_{i}", height=100)
                subj_data[i].update({"优点": adv, "不足": weak, "考情分析": exam, "下一阶段建议": suggest})

        st.divider()

        # --- 生成按钮 ---
        col_btn, col_save = st.columns([2, 1])
        with col_btn:
            if st.button("🚀 生成并下载", type="primary", key="gen_now") and st_name:
                doc = build_archive_doc(st_name, {"学校": school, "年级班级": grade, "入学时间": enroll}, subj_data)
                safe = st_name.replace("/","_")
                buf = io.BytesIO(); doc.save(buf); buf.seek(0)
                st.download_button("📥 下载档案", buf.read(), f"{safe}.docx",
                                   "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                   key=f"dl_now_{safe}")
                st.success(f"✅ {st_name} 的档案已生成")
                st.session_state.archive_students[st_name] = {
                    "info": {"学校": school, "年级班级": grade, "入学时间": enroll},
                    "subjects": copy.deepcopy(subj_data)
                }
                if f"subj_count_{st_name}" in st.session_state:
                    del st.session_state[f"subj_count_{st_name}"]
        with col_save:
            draft = {"school": school, "grade": grade, "enroll": enroll,
                     "subj_count": subj_count, "subj_data": subj_data,
                     "_student_name": st_name}
            st.download_button("💾 保存草稿", json.dumps(draft, ensure_ascii=False, indent=2),
                               f"{st_name or 'draft'}_草稿.json", "application/json",
                               key="save_draft_btn")
            doc = build_archive_doc(st_name, {"学校": school, "年级班级": grade, "入学时间": enroll}, subj_data)
            safe = st_name.replace("/","_")
            buf = io.BytesIO(); doc.save(buf); buf.seek(0)
            st.download_button("📥 下载档案", buf.read(), f"{safe}.docx",
                               "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                               key=f"dl_now_{safe}")
            st.success(f"✅ {st_name} 的档案已生成")
            # 同时保存到批量列表
            st.session_state.archive_students[st_name] = {
                "info": {"学校": school, "年级班级": grade, "入学时间": enroll},
                "subjects": copy.deepcopy(subj_data)
            }
            if f"subj_count_{st_name}" in st.session_state:
                del st.session_state[f"subj_count_{st_name}"]

        # --- 待生成列表 ---
        if st.session_state.archive_students:
            st.divider()
            st.subheader(f"📋 待生成列表（{len(st.session_state.archive_students)} 名学生）")
            for idx, (name, data) in enumerate(st.session_state.archive_students.items()):
                subj_list = ", ".join([s["科目"] for s in data["subjects"]])
                c_info, c_dl = st.columns([5, 1])
                with c_info:
                    st.info(f"**{name}** · {data['info']['学校']} {data['info']['年级班级']} · 科目: {subj_list}")
                with c_dl:
                    doc = build_archive_doc(name, data["info"], data["subjects"])
                    safe = name.replace("/","_")
                    buf = io.BytesIO()
                    doc.save(buf)
                    buf.seek(0)
                    st.download_button(f"📥 {name}", buf.read(), f"{safe}.docx",
                                       "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       key=f"dl_{safe}")

            st.divider()
            col_gen, col_clear = st.columns([1, 3])
            with col_gen:
                if st.button("🚀 打包下载全部（ZIP）", type="primary", key="gen_online"):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        zip_path = os.path.join(tmpdir, "档案汇总.zip")
                        with zipfile.ZipFile(zip_path, 'w') as zf:
                            for name, data in st.session_state.archive_students.items():
                                doc = build_archive_doc(name, data["info"], data["subjects"])
                                safe = name.replace("/","_")
                                dp = os.path.join(tmpdir, f"{safe}.docx")
                                doc.save(dp); zf.write(dp, f"{safe}.docx")
                        with open(zip_path, "rb") as f:
                            st.download_button("📦 下载全部档案（ZIP）", f.read(),
                                               "学生档案汇总.zip", "application/zip")
                        st.success(f"✅ 已生成 {len(st.session_state.archive_students)} 份档案")
            with col_clear:
                if st.button("🗑 清空列表", key="clear_list"):
                    st.session_state.archive_students = {}
                    st.rerun()

# ============================================================
# Tab 2: 课后反馈生成
# ============================================================
with tab2:
    st.header("课后反馈生成")

    mode_fb = st.radio("选择输入方式", ["✏️ 在线填写", "📊 上传Excel"], horizontal=True, key="fb_mode")

    if mode_fb == "📊 上传Excel":
        fb_file = st.file_uploader("课后反馈数据 Excel", type=["xlsx"], key="fb_xlsx")
        if fb_file:
            try:
                wb = openpyxl.load_workbook(fb_file)
                ws1 = wb["上课内容"]
                date = str(ws1.cell(row=3,column=2).value or "").strip()
                lesson = str(ws1.cell(row=4,column=2).value or "").strip()
                ws2 = wb["学生表现"]
                students_fb = []
                for row in ws2.iter_rows(min_row=2, values_only=True):
                    n = str(row[0]).strip() if row[0] else ""
                    nt = str(row[1]).strip() if row[1] else ""
                    if n and nt: students_fb.append({"姓名": n, "表现": nt})
                if st.button("🚀 生成全部反馈", type="primary", key="gen_fb_xlsx"):
                    with tempfile.TemporaryDirectory() as tmpdir:
                        zip_path = os.path.join(tmpdir, "反馈汇总.zip")
                        with zipfile.ZipFile(zip_path, 'w') as zf:
                            all_txt = ""
                            for s in students_fb:
                                txt = f"【课后反馈】{date}\n\n📚 上节课主要内容：\n\n{lesson}\n\n👦 {s['姓名']}课堂表现：\n\n{s['表现']}"
                                safe = s["姓名"].replace("/","_")
                                tp = os.path.join(tmpdir, f"{safe}.txt")
                                with open(tp, "w", encoding="utf-8") as tf: tf.write(txt)
                                zf.write(tp, f"{safe}.txt")
                                all_txt += txt + "\n\n" + "—"*20 + "\n\n"
                            sp = os.path.join(tmpdir, "全部反馈汇总.txt")
                            with open(sp, "w", encoding="utf-8") as sf: sf.write(all_txt)
                            zf.write(sp, "全部反馈汇总.txt")
                        with open(zip_path, "rb") as f:
                            st.download_button("📥 下载全部反馈（ZIP）", f.read(), "课后反馈汇总.zip", "application/zip")
                        st.success(f"✅ 已生成 {len(students_fb)} 份反馈")
            except Exception as e:
                st.error(f"处理失败：{e}")

    else:
        # ===== 在线填写模式 =====
        fb_draft_file = st.file_uploader("📂 恢复草稿", type=["json"], key="fb_draft_load")
        if fb_draft_file:
            st.session_state._fb_draft = json.loads(fb_draft_file.read())
        fb_draft = st.session_state.get("_fb_draft", {})

        st.subheader("📚 上课内容")
        # AI总结功能
        ai_col1, ai_col2 = st.columns([2, 1])
        with ai_col1:
            fb_date = st.text_input("日期", value=fb_draft.get("date","2026.6.10"), key="fb_date")
            fb_course = st.text_input("课程名称", value=fb_draft.get("course",""), placeholder="例：七上U5 A Healthy Lifestyle", key="fb_course")
        with ai_col2:
            lesson_file = st.file_uploader("📎 上传教案自动总结", type=["pdf","docx","txt"], key="lesson_file",
                                           help="支持PDF/Word/文本，自动提取内容并用AI总结为300字")
            if lesson_file:
                file_text = extract_text(lesson_file.read(), lesson_file.name)
                if file_text and not file_text.startswith("["):
                    st.caption(f"已提取 {len(file_text)} 字")
                    api_key = st.session_state.get("deepseek_key", "")
                    if api_key and st.button("✨ AI总结为300字", key="ai_summarize"):
                        with st.spinner("AI总结中…"):
                            summary, err = ai_summarize(file_text, api_key)
                            if err:
                                st.error(err)
                            else:
                                st.session_state._ai_summary = summary
                                st.success("✅ 已生成，已填入下方文本框")
                                st.rerun()
                    elif not api_key:
                        st.caption("⚠️ 请在侧边栏输入DeepSeek API Key")
                else:
                    st.error(file_text)

        ai_summary = st.session_state.get("_ai_summary", "")
        fb_lesson = st.text_area("上课内容（约300字）",
            value=ai_summary or fb_draft.get("lesson",""),
            height=200, key="fb_lesson",
            placeholder="本节课围绕…展开，重点处理了三块内容：\n一是…\n二是…\n三是…")

        st.divider()
        st.subheader("👦 学生表现")

        if "fb_student_count" not in st.session_state:
            draft_count = len(fb_draft.get("students", []))
            st.session_state.fb_student_count = max(draft_count, 3)
        fb_draft_students = fb_draft.get("students", [])

        fb_data = []
        for i in range(st.session_state.fb_student_count):
            dd = fb_draft_students[i] if i < len(fb_draft_students) else {}
            with st.expander(f"学生 {i+1}", expanded=(i==0)):
                c1, c2 = st.columns([1, 4])
                with c1:
                    sname = st.text_input("姓名", value=dd.get("name",""), key=f"fbn_{i}")
                with c2:
                    snotes = st.text_area("今日表现", value=dd.get("notes",""), key=f"fbp_{i}", height=80,
                        placeholder="态度还行，错题重做错误偏多…")
                if sname and snotes:
                    fb_data.append({"姓名": sname, "表现": snotes})

        col_add, col_del, _ = st.columns([1, 1, 4])
        with col_add:
            if st.button("➕ 学生", key="add_fb"):
                st.session_state.fb_student_count += 1; st.rerun()
        with col_del:
            if st.button("➖ 学生", key="del_fb") and st.session_state.fb_student_count > 1:
                st.session_state.fb_student_count -= 1; st.rerun()

        st.divider()

        col_gen, col_save = st.columns([2, 1])
        with col_gen:
            btn_disabled = not (fb_lesson and fb_data)
            if st.button("🚀 生成全部反馈", type="primary", key="gen_fb_online", disabled=btn_disabled):
                if fb_lesson and fb_data:
                    st.success(f"✅ 已生成 {len(fb_data)} 份反馈")
                    for i, s in enumerate(fb_data):
                        txt = f"【课后反馈】{fb_date}\n\n📚 上节课主要内容：\n\n{fb_lesson}\n\n👦 {s['姓名']}课堂表现：\n\n{s['表现']}"
                        with st.expander(f"📄 {s['姓名']}的反馈", expanded=(i==0)):
                            st.text(txt)
                            safe = s["姓名"].replace("/","_")
                            st.download_button(f"📥 下载 {s['姓名']}.txt", txt, f"{safe}.txt", "text/plain", key=f"dl_fb_{i}")
        with col_save:
                fb_draft_out = {"date": fb_date, "course": fb_course, "lesson": fb_lesson,
                                "students": [{"name": s["姓名"], "notes": s["表现"]} for s in fb_data]}
                st.download_button("💾 保存草稿", json.dumps(fb_draft_out, ensure_ascii=False, indent=2),
                                   f"{fb_date}_反馈草稿.json", "application/json", key="save_fb_draft")
# ============================================================
# 侧边栏
# ============================================================
with st.sidebar:
    st.markdown("### 🔑 DeepSeek API Key")
    api_key = st.text_input("输入Key解锁AI总结功能", type="password", key="deepseek_key_input",
                            help="在 platform.deepseek.com 获取")
    if api_key:
        st.session_state.deepseek_key = api_key
        st.success("已设置")

    st.divider()
    st.markdown("### 📖 使用说明")
    st.markdown("""
    **两种输入方式，切换顶部开关即可：**

    ✏️ **在线填写**（推荐）
    - 直接在网页上填，不用Excel
    - 档案生成：选学生→填各科内容→加入列表→一键生成
    - 反馈生成：填上课内容+学生表现→生成

    📊 **上传Excel**
    - 兼容旧版Excel模板
    - 适用于已有Excel数据的场景

    ---
    **输出格式：**
    - 学生档案 → .docx（Word文档）
    - 课后反馈 → .txt（复制发微信）
    """)
